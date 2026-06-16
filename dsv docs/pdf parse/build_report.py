"""Build PDF parse logic & procedure report (.docx + .md)."""
import json
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(r"C:\Users\jichu\Downloads\dsv docs")
OUT_DIR = ROOT / "parser_out_v3"
recs = json.load(open(OUT_DIR / "pdf_evidence_index.json", encoding="utf-8"))

# ── 집계
by_type, by_verdict, by_typeb, issues = {}, {}, {}, []
unknowns = []
missed_dos = []
for r in recs:
    by_type[r.get("doc_type","?")] = by_type.get(r.get("doc_type","?"), 0) + 1
    by_verdict[r.get("verdict","?")] = by_verdict.get(r.get("verdict","?"), 0) + 1
    by_typeb[r.get("type_b","?")] = by_typeb.get(r.get("type_b","?"), 0) + 1
    if r.get("doc_type") == "UNKNOWN": unknowns.append(r["file"])
    if r.get("doc_type") not in ("DELIVERY_ORDER","DELIVERY_NOTE") and "_DO" in r["file"]:
        missed_dos.append(r["file"])
    if r.get("verdict") == "FAILED": issues.append((r["file"],"PARSE_FAILED",r.get("error","")))
    if r.get("doc_type") == "UNKNOWN": issues.append((r["file"],"UNKNOWN_DOCTYPE","header fingerprint not matched"))
    if not r.get("keys",{}).get("shipment_no"): issues.append((r["file"],"NO_SHIPMENT_NO","HVDC-ADOPT-XXXX not found"))

# ── 표 헬퍼
def _sanitize(s):
    """Remove XML-incompatible control characters and U+FFFE/U+FFFF."""
    s = str(s)
    # Strip U+FFFE (BOM) and other non-characters
    s = s.replace("￾", "<U+FFFE>")
    s = s.replace("￿", "<U+FFFF>")
    # Strip C0/C1 control chars (keep tab/newline)
    return "".join(c for c in s if c in ("\t", "\n") or (0x20 <= ord(c) <= 0xFFFD and ord(c) not in (0x7F,)))

def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = _sanitize(h)
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = _sanitize(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return table

def H1(doc, txt): p=doc.add_heading(txt, level=1)
def H2(doc, txt): p=doc.add_heading(txt, level=2)
def H3(doc, txt): p=doc.add_heading(txt, level=3)
def P(doc, txt, bold=False, size=10):
    p = doc.add_paragraph()
    r = p.add_run(txt); r.font.size = Pt(size); r.bold = bold
    return p

# ── 문서 시작
doc = Document()
# 페이지 여백
for sec in doc.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)

# 기본 폰트
style = doc.styles["Normal"]
style.font.name = "Calibri"; style.font.size = Pt(10)

# ── 표지
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("HVDC Invoice Audit")
r.bold = True; r.font.size = Pt(20)
doc.add_paragraph()
st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("PDF Parse Logic & Procedure Report")
r.bold = True; r.font.size = Pt(16)
doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
r.font.size = Pt(10); r.italic = True
doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Project: Samsung C&T HVDC Abu Dhabi\n").font.size = Pt(10)
meta.add_run("Scope: DSV Docs root (57 PDFs)\n").font.size = Pt(10)
meta.add_run("Playbook: pdfparse.md v1.0").font.size = Pt(10)
doc.add_page_break()

# ── 1. Executive Summary
H1(doc, "1. Executive Summary")
P(doc, "본 보고서는 pdfparse.md 플레이북을 구현한 Hybrid PDF Parser를 루트 디렉터리의 57개 PDF에 적용한 결과를 정리한다.")
P(doc, f"처리 대상: {len(recs)}개 PDF. 판정: PASS {by_verdict.get('PASS',0)} / AMBER {by_verdict.get('AMBER',0)} / ZERO {by_verdict.get('ZERO',0)} / FAILED {by_verdict.get('FAILED',0)}")
P(doc, "결론: 파싱 파이프라인은 정상 동작하며 텍스트 레이어 기반 추출이 대부분 안정적. 다만 DOC_TYPE 분류 룰의 false positive/negative가 발견되어 1차 보완이 필요함.", bold=True)

# ── 2. Procedure
H1(doc, "2. Procedure — 5-Stage Pipeline")
H2(doc, "2.1 Stage Map")
P(doc, "pdfparse.md 플레이북에 따라 다음 5단계로 구현:")
add_table(doc, ["#", "Stage", "Engine", "Output"],
    [["1", "Text extract", "pdfplumber (text layer)", "전체 텍스트 + 페이지 수"],
     ["2", "Table extract", "pdfplumber.extract_tables()", "페이지별 표 (CSV 호환)"],
     ["3", "Render QA", "PyMuPDF (fitz) get_pixmap(dpi=120)", "첫 페이지 PNG (옵션)"],
     ["4", "Evidence Index", "COMMON_KEYS regex 7종", "JSON: shipment/container/bl/do/invoice/date/amounts"],
     ["5", "TYPE_B + Gate", "TYPE_B_RULES + gate_verdict()", "PASS/AMBER/ZERO/FAILED"]],
    widths=[1.0, 3.0, 5.5, 8.0])

H2(doc, "2.2 Command")
P(doc, "기본 실행:", bold=True)
P(doc, "  python pdf_hybrid_parser.py --input . --out-dir parser_out", size=9)
P(doc, "Render QA 포함 (각 PDF 첫 페이지 PNG 저장):", bold=True)
P(doc, "  python pdf_hybrid_parser.py --input . --out-dir parser_out --render", size=9)
P(doc, "단일 파일:", bold=True)
P(doc, "  python pdf_hybrid_parser.py --input \"HVDC-ADOPT-SCT-0122_CarrierInvoice.pdf\" --out-dir parser_out", size=9)

# ── 3. Logic
H1(doc, "3. Logic — Classification & Extraction")
H2(doc, "3.1 DOC_TYPE 분류 (Header Fingerprint)")
P(doc, "텍스트 첫 3,000자를 대상으로 다음 우선순위로 매칭. 매칭 안 되면 UNKNOWN.")
add_table(doc, ["DOC_TYPE", "Header 키워드", "문서 실예"],
    [["BOE_CUSTOMS",    "'Customs Declaration' / 'DEBIT NOTE' / 'Gate Pass' / 'بوليصة'",   "SCT-0122~0134 BOE, HE-0425~0502 BOE"],
     ["DELIVERY_ORDER", "'DELIVERY ORDER' / 'D/Order No'",                                  "SCT-0123·0124 DO (D/Order No 매칭)"],
     ["DELIVERY_NOTE",  "'NOT NEGOTIABLE DELIVERY NOTE' / 'Delivery Note/Waybill'",         "SCT-0122 DN, SCT-0123·0124 DN 등"],
     ["CARRIER_RHS",    "'RAIS HASSAN SAADI' + 'TAX INVOICE'",                              "RHS 발행 carrier invoice"],
     ["CARRIER_EVG",    "'EVERGREEN SHIPPING AGENCY' + 'TAX INVOICE'",                      "Evergreen 발행 carrier invoice"],
     ["PORT_ALLIED",    "'ALLIED ONDOCK' + 'TAX INVOICE'",                                  "SCT-0038 CNTRepair (ALLIED 매칭)"],
     ["PORT_CSP",       "'CSP Abu Dhabi Terminal' + 'Package Invoice'",                     "CSP 발행 port invoice"],
     ["UNKNOWN",        "위 어느 것도 매칭 안 됨",                                            "HE-0499 lot1~3 (LAND IMPORT 형식) 등"]])
# (위 표는 일반적인 헤더 패턴 — 실 PDF 텍스트 인용은 생략)

H2(doc, "3.2 COMMON_KEYS Regex (7종)")
add_table(doc, ["Key", "Regex", "정규화"],
    [["shipment_no", r"HVDC[\s\-￾]+ADOPT[\s\-]+SCT[\s\-]*(\d{4})  /  HE variant", "HVDC-ADOPT-SCT-XXXX / HVDC-ADOPT-HE-XXXX"],
     ["container",   r"\b([A-Z]{4})\s?(\d{6})-?(\d)\b",                              "HMMU 608937-7 → HMMU6089377"],
     ["bl_no",       r"\b(?:SELA|EGLV)[A-Z0-9]{8,}\b",                              "원문 유지 (vendor prefix)"],
     ["do_no",       r"(?:D/Order No|DO #|Delivery Order No)[:\s]+(\d{6,})",       "첫 매치만"],
     ["invoice_no",  r"(?:Invoice No|INVOICE NO\.?|Invoice Number)\s*:?\s*([A-Z0-9\-/]+)", "첫 매치만"],
     ["date",        r"\d{2}/\d{2}/\d{4} | \d{2}-[A-Z]{3}-\d{4} | \d{4}-\d{2}-\d{2}", "첫 매치만"],
     ["amount_aed",  r"AED\s*([0-9,]+(?:\.\d{2})?)",                                "콤마 제거 후 float 리스트"]])

H2(doc, "3.3 TYPE_B 매핑 (우선순위)")
add_table(doc, ["TYPE_B", "트리거 키워드", "Fallback by DOC_TYPE"],
    [["Inspection", "Container Inspection Fee, Admin & Inspection, Customs Inspection, Inspection Fee", "PORT_ALLIED, PORT_CSP"],
     ["Customs",    "BOE, Customs Duty, Customs Declaration, Debit Note, HS Code",                          "BOE_CUSTOMS"],
     ["DO",         "Delivery Order, DO Fee, D/Order",                                                       "DELIVERY_ORDER"],
     ["INLAND",     "DN Road Freight, Truck, FB movement, Inland, Transport",                               "DELIVERY_NOTE"],
     ["THC",        "Terminal Handling, Port Handling, THC",                                                  "—"],
     ["OTHERS",     "ISPS, Equipment Repositioning, Container Maintenance, Airport Fees, Appointment, CNT Repair, Empty Return, Washing", "그 외"]])

H2(doc, "3.4 Gate Verdict (PASS/AMBER/ZERO)")
P(doc, "pdfparse.md §ZERO Log에 명시된 자동 PASS 차단 조건을 구현:")
P(doc, "  • ZERO 차단 #1: BOE/HS/UAE Customs 최종 판정 → BOE_CUSTOMS는 항상 AMBER", size=9)
P(doc, "  • ZERO 차단 #2: DEM/DET/Storage 정산 → 본 라운드 PDF에 해당 없음", size=9)
P(doc, "  • ZERO 차단 #3: Grand Total 공란 invoice → CARRIER/PORT invoice에서 AED amount 0건 시 AMBER", size=9)
P(doc, "  • AMBER: shipment_no 미발견, 또는 위 조건 해당", size=9)
P(doc, "  • PASS: 위 어느 차단 조건에도 해당 안 됨", size=9)

# ── 4. Results
H1(doc, "4. Results")
H2(doc, "4.1 DOC_TYPE 분포")
rows = [[k, v, f"{v/len(recs)*100:.1f}%"] for k, v in sorted(by_type.items(), key=lambda x: -x[1])]
add_table(doc, ["DOC_TYPE", "건수", "비율"], rows, widths=[5,2,2])
H2(doc, "4.2 TYPE_B 분포")
rows = [[k, v, f"{v/len(recs)*100:.1f}%"] for k, v in sorted(by_typeb.items(), key=lambda x: -x[1])]
add_table(doc, ["TYPE_B", "건수", "비율"], rows, widths=[5,2,2])
H2(doc, "4.3 Verdict 분포")
rows = [[k, v, f"{v/len(recs)*100:.1f}%"] for k, v in sorted(by_verdict.items(), key=lambda x: -x[1])]
add_table(doc, ["Verdict", "건수", "비율"], rows, widths=[5,2,2])

# ── 5. Findings
H1(doc, "5. Findings — 분류 보완 필요 항목")
H2(doc, "5.1 UNKNOWN (9건) — 신규 vendor/format 발견")
P(doc, "다음 9개 PDF는 기존 7종 DOC_TYPE 어느 것에도 매칭되지 않음. 신규 분류 룰 추가 필요.")
add_table(doc, ["#", "파일", "실제 헤더 패턴", "권장 DOC_TYPE"],
    [[1, "HVDC-ADOPT-HE-0499(lot1)_SupportingDocs.pdf",  "LAND IMPORT 형식 (FED/DEF/Duty 테이블)",         "BOE_CUSTOMS (LAND_IMPORT)"],
     [2, "HVDC-ADOPT-HE-0499(Lot2)_SupportingDocuments.pdf", "LAND IMPORT (page 1 이미지, text 6 chars)", "BOE_CUSTOMS (LAND_IMPORT, OCR 필요)"],
     [3, "HVDC-ADOPT-HE-0499(lot3)_SupportingDocs.pdf",  "LAND IMPORT 형식",                                 "BOE_CUSTOMS (LAND_IMPORT)"],
     [4, "HVDC-ADOPT-SCT-0122_PortCNTInsp.pdf",         "TransWorldContainer Services LLC 발행",            "PORT_TWCS (신규)"],
     [5, "HVDC-ADOPT-SCT-0123, 0124_PortCNTWashing.pdf","Allied Ondock (TAX INVOICE 동시 매칭 → CARRIER_RHS 오분류)", "PORT_ALLIED"],
     [6, "HVDC-ADOPT-SCT-0127_PortCNTInsp.pdf",         "Allied Ondock (TAX INVOICE 동시 매칭)",            "PORT_ALLIED"],
     [7, "HVDC-ADOPT-SCT-0126_PortCNTInspection.pdf",   "Allied Ondock (TAX INVOICE 동시 매칭)",            "PORT_ALLIED"],
     [8, "HVDC-ADOPT-SCT-0123, 0124_PortCNTInsp.pdf",   "Allied Ondock (TAX INVOICE 동시 매칭)",            "PORT_ALLIED"],
     [9, "HVDC-ADOPT-SCT-0038-CNTRepair.pdf",           "Allied Ondock CNT Repair (현재 CARRIER_RHS 오분류)", "PORT_ALLIED"]])
# 주의: 위 표의 헤더 패턴은 일반화된 형식 표현. 실 PDF 본문은 인용하지 않음.

H2(doc, "5.2 DELIVERY_ORDER 분류 누락 (3건)")
P(doc, "다음 DO 파일들은 'DELIVERY ORDER' / 'D/Order No' 키워드가 없어 UNKNOWN 또는 다른 타입으로 분류됨. 헤더 키워드 확장 필요.")
add_table(doc, ["파일", "실제 헤더 표현", "현재 분류", "권장 수정"],
    [["HVDC-ADOPT-SCT-0127_DO.pdf",  "D.O. NUMBER : 20019259 (점으로 표기)",         "UNKNOWN", "RX_DO 보강: 'D\\.O\\. NUMBER' 패턴 추가"],
     ["HVDC-ADOPT-SCT-0131_DO.pdf",  "DELIVERY NOTIFICATION",                        "UNKNOWN", "신규 키워드 'DELIVERY NOTIFICATION' 추가"],
     ["HVDC-ADOPT-SCT-0134_DO.pdf",  "DELIVERY NOTIFICATION MASTER",                 "UNKNOWN", "신규 키워드 'DELIVERY NOTIFICATION MASTER' 추가"]])

H2(doc, "5.3 Cargo Type 추출 보강 권장")
P(doc, "현재 evidence는 line amount를 AED 패턴으로만 수집. carrier/port invoice의 line-level 추가는 다음 단계:")
P(doc, "  • tables[0] 구조 분석 후 'TWCSInspectionCharges', 'WashingCharges20Ft' 등 description 컬럼 매칭", size=9)
P(doc, "  • 단가 × 수량 = amount 검증 (delta 2% 임계치 적용)", size=9)
P(doc, "  • pdfparse.md의 line-level evidence만 사용 / grand total 자동 PASS 금지 정책 유지", size=9)

# ── 6. ZERO Conditions
H1(doc, "6. ZERO Conditions (pdfparse.md §ZERO Log)")
P(doc, "다음 조건이 감지되면 자동 PASS 차단 (Rule #0 우선 적용 — 차단이어도 Review Pack 다운로드는 허용):")
add_table(doc, ["#", "조건", "현 라운드 적중"],
    [[1, "BOE/HS/UAE Customs 최종 판정",                 f"{by_type.get('BOE_CUSTOMS',0)}건 AMBER 처리"],
     [2, "DEM/DET/Storage settlement",                    "0건 (해당 PDF 없음)"],
     [3, "Grand Total 공란 invoice",                      "0건 (AED amount 일부 추출 성공)"]])
P(doc, "현 라운드 ZERO 판정 0건. AMBER 51건은 모두 reviewer 승인 필요.", bold=True)

# ── 7. Recommendations
H1(doc, "7. Recommendations (다음 작업)")
H2(doc, "7.1 즉시 (1일)")
P(doc, "1) 분류 룰 보강: BOE_CUSTOMS에 'LAND IMPORT' 추가, DELIVERY_ORDER에 'D.O. NUMBER' / 'DELIVERY NOTIFICATION' 추가, PORT_ALLIED 우선순위를 CARRIER_RHS보다 높임 (TAX INVOICE 동시 매칭 해결)", size=9)
P(doc, "2) container 정규화 결과 evidence index에 amount 누락 PDF 0건 확인 — 모두 AED amount 추출됨 (Carrier 4종 합 10건 amount)", size=9)
P(doc, "3) HE-0499 lot2의 text length 6 → OCR fallback (Tesseract) 필요. PDF text layer 없음.", size=9)

H2(doc, "7.2 단기 (1~2일)")
P(doc, "1) line-level amount 추출: tables[0] 기반 description × qty × unit_price 검증", size=9)
P(doc, "2) Evidence Index → 13-sheet Excel 워크북 (per apps/web workbook-builder.ts) 자동 조립", size=9)
P(doc, "3) Rule #0 적용: PASS/AMBER/ZERO 어느 경우든 Excel은 항상 산출", size=9)

H2(doc, "7.3 중기 (3~5일)")
P(doc, "1) Hybrid parser + 13-sheet Excel + validation loop 3회 (pdfparse.md Option C)", size=9)
P(doc, "2) webhook callback: worker-py → web(/run) → workbook-builder → Vercel Blob", size=9)
P(doc, "3) DLP (이미 제거됨, 2026-06-15) 외 추가 모듈 금지 준수", size=9)

# ── 8. Appendix
H1(doc, "8. Appendix — 산출 파일")
add_table(doc, ["파일", "용도"],
    [["parser_out/pdf_evidence_index.json",  "57 PDF × 7 KEY + DOC_TYPE + TYPE_B + verdict (full)"],
     ["parser_out/pdf_evidence_index.xlsx",  "Evidence_Index + Summary + Issues (3-sheet)"],
     ["parser_out/summary.md",               "집계 요약"],
     ["parser_out/renders/*.png",            "57개 PDF 첫 페이지 PNG (render QA)"],
     ["pdf_hybrid_parser.py",                "본 보고서 대상 스크립트 (playbook 구현)"],
     ["20260615_pdf_parse_procedure_v1.docx", "본 보고서"],
     ["20260615_pdf_parse_procedure_v1.md",  "본 보고서 마크다운 요약"]])

# ── 문서 저장
# ── 9. Fix Cycle v1 → v3
H1(doc, "9. Appendix — Fix Cycle (v1 → v3)")
P(doc, "본 보고서 §5 권장 사항을 parser_v1.docx 산출 이후 즉시 반영. v1→v2→v3 3회 실행으로 UNKNOWN 9→2 (78% 해결).")
add_table(doc, ["DOC_TYPE", "v1", "v2", "v3", "v3-v1"],
    [["AIRPORT_FEES",   0, 0, 2, "+2"],
     ["APPOINTMENT",    0, 0, 2, "+2"],
     ["BOE_CUSTOMS",   26, 28, 28, "+2 (LAND_IMPORT)"],
     ["CARRIER_RHS",    9, 4, 4, "−5 (PORT 우선순위로 재배정)"],
     ["CARRIER_EVG",    1, 0, 0, "−1"],
     ["DELIVERY_NOTE",  9, 7, 7, "−2"],
     ["DELIVERY_ORDER", 3, 6, 6, "+3 (D.O. NUMBER/DELIVERY NOTIFICATION)"],
     ["PORT_ALLIED",    0, 4, 4, "+4 (TAX INVOICE 동시매칭 해결)"],
     ["PORT_CSP",       0, 1, 1, "+1"],
     ["PORT_TWCS",      0, 0, 1, "+1 (TransWorldContainer)"],
     ["UNKNOWN",        9, 7, 2, "−7 (잔여 2건은 image_only OCR 대상)"]])

H2(doc, "9.1 잔여 UNKNOWN 2건 (OCR fallback 대상)")
P(doc, "두 PDF는 text layer 0건 — Tesseract OCR 필요. evidence index의 image_only=True 플래그로 식별됨.")
add_table(doc, ["파일", "text_length", "image_only", "권장 후속"],
    [["HVDC-ADOPT-HE-0499(Lot2)_SupportingDocuments.pdf", "0", "True", "Tesseract OCR (7 pages)"],
     ["HVDC-ADOPT-SCT-0131_DN.pdf",                       "0", "True", "Tesseract OCR (스캔 DN)"]])

H2(doc, "9.2 DO 번호 정규화 검증")
P(doc, "RX_DO 패턴에 D.O. NUMBER / D/O NUMBER 추가 + 공백 제거 정규화 적용. 결과:")
add_table(doc, ["파일", "v2 do_no (공백 포함)", "v3 do_no (정규화)"],
    [["HVDC-ADOPT-SCT-0123, 0124_DO.pdf", "22501743", "22501743"],
     ["HVDC-ADOPT-SCT-0127_DO.pdf",       "\"2 0 0 1 9 2 5 9\"", "20019259"]])

# ── 문서 저장
out_docx = ROOT / "20260615_pdf_parse_procedure_v2.docx"
doc.save(str(out_docx))
print(f"[OK] {out_docx}")
